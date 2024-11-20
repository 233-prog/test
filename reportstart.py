import yfinance as yf
from docx import Document

# Define the stock ticker for TBZ (Tribhovandas Bhimji Zaveri Ltd)
ticker_symbol = "TBZ.NS"
stock = yf.Ticker(ticker_symbol)

# Fetch the most recent stock data and general info
data = stock.history(period="1y")
info = stock.info

# Get stock details
current_price = data['Close'][-1]
market_cap = info.get("marketCap")
year_high = data['High'].max()
year_low = data['Low'].min()
fifty_day_ma = data['Close'].rolling(window=50).mean().iloc[-1]
two_hundred_day_ma = data['Close'].rolling(window=200).mean().iloc[-1]
one_year_change = (current_price - data['Close'][0]) / data['Close'][0] * 100

# Calculate RSI
delta = data['Close'].diff()
gain = delta.where(delta > 0, 0)
loss = -delta.where(delta < 0, 0)
avg_gain = gain.rolling(window=14).mean().iloc[-1]
avg_loss = loss.rolling(window=14).mean().iloc[-1]
rs = avg_gain / avg_loss
rsi = 100 - (100 / (1 + rs))

# position calc

position=(current_price - year_low) / (year_high - year_low) * 100
position_text= ""
if position >= 80:
    position_text= "The stock is in a strong uptrend, near its 1-year high."
elif 60 <= position < 80:
    position_text= "The stock shows a moderate uptrend, trending upward but not at its peak."
elif 40 <= position < 60:
    position_text= "The stock price is stable, trading around the middle of its 1-year range."
elif 20 <= position < 40:
    position_text= "The stock shows a moderate downtrend, below the midpoint of its 1-year range."
else:
    position_text= "The stock is in a strong downtrend, near its 1-year low."

print(current_price, year_low, year_high)
print(position)
print(position_text)
exit()

# Print report
print(f"Tribhovandas Bhimji Zaveri Ltd. (TBZ)")
print(f"Date: {data.index[-1].date()}")
print(f"Ticker: {ticker_symbol}")
print(f"Sector: Jewellery and Retail")
print(f"Exchange: NSE/BSE")
print(f"Market Capitalization: ₹{market_cap / 1e7:.2f} crore (approx.)")
print(f"Current Stock Price: ₹{current_price:.2f} (as of recent closing)")

print("\nPrice History and Performance")
print(f"1-Year High: ₹{year_high:.2f}")
print(f"1-Year Low: ₹{year_low:.2f}")
print(f"52-Week Change: {one_year_change:.2f}%")

print("\nTechnical Indicators")
print(f"Relative Strength Index (RSI): {rsi:.2f}")
print(f"50-Day Moving Average: ₹{fifty_day_ma:.2f}")
print(f"200-Day Moving Average: ₹{two_hundred_day_ma:.2f}")
print(f"The price trend: {position}")

doc = Document()


doc.add_paragraph(f"Tribhovandas Bhimji Zaveri Ltd. (TBZ)")
doc.add_paragraph(f"200-Day Moving Average: ₹{two_hundred_day_ma:.2f}")

# Save the document
output_filename = f"{ticker_symbol}_Stock_Info.docx"
doc.save(output_filename)

print(f"Stock information has been saved to {output_filename}.")